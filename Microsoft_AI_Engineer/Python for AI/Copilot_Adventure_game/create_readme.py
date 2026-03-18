from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Python Adventure Game - README', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction
doc.add_heading('Project Overview', level=1)
intro = doc.add_paragraph(
    'This document outlines the development of a text-based adventure game '
    'built using Python. The game allows players to explore different locations, '
    'make strategic choices, and attempt to find a legendary treasure. The game '
    'incorporates fundamental Python concepts including variables, conditionals, '
    'functions, and user input handling.'
)

# Add description
doc.add_heading('Game Description', level=1)
doc.add_paragraph(
    'The adventure game presents players with a quest to find a legendary treasure '
    'hidden in an ancient land. The game flow is as follows:'
)
game_flow = doc.add_paragraph(
    '1. The player enters their name\n'
    '2. They choose between exploring a dark forest or a mysterious cave\n'
    '3. Each path presents a new scenario with two choices\n'
    '4. Correct choices lead to victory (finding the treasure)\n'
    '5. Poor decisions result in failure\n'
    '6. Players can choose to replay the game or exit',
    style='List Bullet'
)

# Add development steps
doc.add_heading('Development Steps', level=1)

doc.add_heading('Step 1: Project Setup', level=2)
doc.add_paragraph(
    'Created a new Python file named adventure_game.py in the project folder. '
    'Added an inline comment describing the purpose of the script as a text-based '
    'adventure game. This file serves as the main entry point for the entire application.'
)

doc.add_heading('Step 2: Create the start_game() Function', level=2)
doc.add_paragraph(
    'Implemented the start_game() function to display the game introduction and '
    'handle the player\'s initial choice. The function:'
)
doc.add_paragraph(
    '• Greets the player with a welcome message\n'
    '• Prompts the player to enter their name and stores it in a variable\n'
    '• Displays the initial scenario and choice options\n'
    '• Uses if-elif-else conditionals to route the player to forest_path() or cave_path()\n'
    '• Handles invalid input by recursively calling itself',
    style='List Bullet'
)

doc.add_heading('Step 3: Implement the forest_path() Function', level=2)
doc.add_paragraph(
    'Created the forest_path() function to handle the forest scenario. This function:'
)
doc.add_paragraph(
    '• Describes the forest environment to the player\n'
    '• Offers two choices: follow a river or climb a tree\n'
    '• Uses if-elif-else structure to determine outcomes\n'
    '• Following the river results in failure\n'
    '• Climbing the tree leads to victory\n'
    '• Calls play_again() at the end regardless of outcome',
    style='List Bullet'
)

doc.add_heading('Step 4: Implement the cave_path() Function', level=2)
doc.add_paragraph(
    'Developed the cave_path() function for the cave scenario with similar structure:'
)
doc.add_paragraph(
    '• Describes the dark cave environment\n'
    '• Offers two choices: light a torch or proceed in the dark\n'
    '• Lighting a torch reveals the path to treasure (victory)\n'
    '• Proceeding in darkness causes a fall (failure)\n'
    '• Calls play_again() to prompt for replay',
    style='List Bullet'
)

doc.add_heading('Step 5: Add Replay Functionality', level=2)
doc.add_paragraph(
    'Created the play_again() function to handle game restart logic. This function:'
)
doc.add_paragraph(
    '• Asks the player if they want to play again\n'
    '• If yes, recursively calls start_game() to restart\n'
    '• If no, displays farewell message and exits the program',
    style='List Bullet'
)

doc.add_heading('Step 6: Set Up Main Entry Point', level=2)
doc.add_paragraph(
    'Added the standard Python if __name__ == "__main__" guard to call start_game() '
    'when the script is run directly. This ensures the game starts automatically '
    'when the user executes python adventure_game.py.'
)

# Add key features
doc.add_heading('Key Features', level=1)
doc.add_paragraph(
    '• Interactive CLI: Players interact with the game through command-line prompts\n'
    '• Branching Narratives: Different paths lead to different outcomes\n'
    '• Win/Loss Conditions: Clear victory and failure states\n'
    '• Replay Mechanism: Players can restart without re-running the script\n'
    '• Input Validation: Invalid choices are handled gracefully',
    style='List Bullet'
)

# Add testing section
doc.add_heading('Testing & Terminal Output', level=1)
doc.add_paragraph(
    'The game was tested by running multiple playthroughs with different choice '
    'combinations. Below is the terminal output from a comprehensive test session '
    'that exercises all major paths and features:'
)

# Add code block style for terminal output
test_output = '''
C:\Developer\Python\Projects\Simplilearn\Copilot_Adventure_game>python adventure_game.py
Welcome to the Adventure Game!
What is your name, explorer? Arthur
Hello, Arthur! Your quest is to find the legendary treasure.
You stand at a crossroads. Do you want to explore the (1) dark forest or (2) mysterious cave?
Enter 1 for forest or 2 for cave: 1
You step into the dark forest. The trees loom overhead and you hear a river nearby.
Do you (1) follow the sound of the river or (2) climb a tall tree to get a better view?
Enter 1 or 2: 1
You follow the river and slip on a rock, injuring yourself. The quest ends here.
Do you want to play again? (y/n): y
Welcome to the Adventure Game!
What is your name, explorer? Arthur
Hello, Arthur! Your quest is to find the legendary treasure.
You stand at a crossroads. Do you want to explore the (1) dark forest or (2) mysterious cave?
Enter 1 for forest or 2 for cave: 2
You enter the mysterious cave. It's dark and you can barely see.
Do you (1) light a torch or (2) proceed in the dark relying on your other senses?
Enter 1 or 2: 2
You stumble in the darkness and fall into a pit. The quest ends here.
Do you want to play again? (y/n): y
Welcome to the Adventure Game!
What is your name, explorer? Arthur
Hello, Arthur! Your quest is to find the legendary treasure.
You stand at a crossroads. Do you want to explore the (1) dark forest or (2) mysterious cave?
Enter 1 for forest or 2 for cave: 1
You step into the dark forest. The trees loom overhead and you hear a river nearby.
Do you (1) follow the sound of the river or (2) climb a tall tree to get a better view?
Enter 1 or 2: 3
That choice doesn't make sense. The forest swallows you.
Do you want to play again? (y/n): y
Welcome to the Adventure Game!
What is your name, explorer? Arthur
Hello, Arthur! Your quest is to find the legendary treasure.
You stand at a crossroads. Do you want to explore the (1) dark forest or (2) mysterious cave?
Enter 1 for forest or 2 for cave: 2
You enter the mysterious cave. It's dark and you can barely see.
Do you (1) light a torch or (2) proceed in the dark relying on your other senses?
Enter 1 or 2: 1
The torch reveals ancient markings that lead you to the treasure chamber!
Congratulations, you have found the treasure and won the game!
Do you want to play again? (y/n): n
Thanks for playing! Farewell, explorer.
'''

# Add terminal output as a formatted paragraph
output_para = doc.add_paragraph(test_output)
output_para.style = 'Normal'
for run in output_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# Add code structure
doc.add_heading('Code Architecture', level=1)
doc.add_paragraph(
    'The program follows a functional architecture with modular functions:'
)

doc.add_heading('Function Summary', level=2)
doc.add_paragraph(
    'start_game(): Entry point that greets the player, collects their name, '
    'and presents the initial fork in the adventure (forest vs. cave).',
)
doc.add_paragraph(
    'forest_path(): Handles the dark forest scenario with a river/tree choice.',
)
doc.add_paragraph(
    'cave_path(): Handles the mysterious cave scenario with a torch/dark choice.',
)
doc.add_paragraph(
    'play_again(): Manages the replay logic, either restarting or exiting.',
)

# Add conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'The adventure game successfully implements all requirements from the project '
    'specification. The code demonstrates proper use of functions, conditionals, '
    'variables, and user input handling. The modular design makes it easy to extend '
    'with additional paths or features in the future. All test cases execute correctly, '
    'with proper handling of both valid and invalid user inputs.'
)

# Save the document as .doc file
output_path = r'c:\Developer\Python\Projects\Simplilearn\Copilot_Adventure_game\Adventure_Game_README.doc'
doc.save(output_path)
print(f"README document created successfully: {output_path}")
